#!/usr/bin/env python3
"""生成 day7-CSRF漏洞修复报告.docx（仅包含CSRF相关漏洞）"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 封面 ──
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CSRF漏洞修复报告')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Web应用安全检测与修复')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph('')
doc.add_paragraph('')

meta_items = [
    ('报告编号', 'SEC-CSRF-2026-001'),
    ('检测目标', '用户管理系统 Web 应用'),
    ('检测日期', '2026-07-24'),
    ('应用版本', '1.0'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_vuln_table(title_text, severity, consequences, fix_method, before_code, after_code):
    """添加每个漏洞的完整描述"""
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 50, 50)

    p = doc.add_paragraph()
    run = p.add_run('【严重程度】')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(f'  {severity}').font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('【不修复的后果】')
    run.font.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(consequences)
    p2.paragraph_format.left_indent = Cm(0.5)
    for r in p2.runs:
        r.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run('【修复方法】')
    run.font.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(fix_method)
    p2.paragraph_format.left_indent = Cm(0.5)
    for r in p2.runs:
        r.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run('【修复前后对比】')
    run.font.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(['修复前', '修复后']):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for i, code in enumerate([before_code, after_code]):
        cell = table.rows[1].cells[i]
        cell.text = ''
        for line in code.split('\n'):
            p = cell.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
        if cell.paragraphs[0].text == '' and len(cell.paragraphs) > 1:
            p_element = cell.paragraphs[0]._element
            p_element.getparent().remove(p_element)

    for row in table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)

    doc.add_paragraph('')

# ======================== 正文开始 ========================

# ── 一、检测结果汇总 ──
add_heading_styled('一、检测结果汇总', level=1)

summary_table = doc.add_table(rows=6, cols=4)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['接口路径', '请求方法', 'CSRF防护状态', '检测结论']
for i, h in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
    cell._tc.get_or_add_tcPr().append(shading)

data = [
    ['/login', 'POST', '已启用（CSRF Token）', '安全'],
    ['/register', 'POST', '已启用（CSRF Token）', '安全'],
    ['/change-password', 'POST', '已启用（CSRF Token）', '已修复'],
    ['/upload', 'POST', '已启用（CSRF Token）', '安全'],
    ['/recharge', 'POST', '已启用（CSRF Token）', '安全'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        cell = summary_table.rows[row_idx].cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if val == '已修复':
            run.font.color.rgb = RGBColor(200, 130, 0)
            run.font.bold = True
        elif val == '安全':
            run.font.color.rgb = RGBColor(0, 120, 0)

col_widths = [Inches(1.8), Inches(0.8), Inches(2.2), Inches(2.2)]
for row in summary_table.rows:
    for i, width in enumerate(col_widths):
        row.cells[i].width = width

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('检测结论：')
run.font.bold = True
run.font.size = Pt(11)
p.add_run(
    '本次检测共发现 1 个 CSRF 漏洞，位于 /change-password 接口。'
    '该接口在处理密码修改请求时未执行任何 CSRF 防护措施。'
    '其余 POST 接口均已启用 CSRF Token 验证机制，不存在 CSRF 风险。'
    '该漏洞已完成修复。'
).font.size = Pt(10.5)

doc.add_page_break()

# ── 二、漏洞详情 ──
add_heading_styled('二、漏洞详情', level=1)

add_vuln_table(
    '漏洞一：/change-password 接口缺少 CSRF 防护',
    '严重（Critical）',
    '攻击者可以构造恶意网页，诱导已登录用户访问。该恶意页面会自动向 /change-password '
    '接口提交伪造的 POST 请求，携带攻击者指定的用户名和新密码。由于该接口未验证 CSRF Token，'
    '请求会被服务器正常处理，导致以下后果：\n'
    '1. 任意用户密码被篡改：攻击者可修改包括管理员在内的任意账户密码，实现账户劫持。\n'
    '2. 权限提升：攻击者通过修改管理员密码获取管理员权限，进一步控制系统。\n'
    '3. 数据泄露：账户被控制后，攻击者可查看所有用户信息及敏感数据。\n'
    '4. 完全接管：结合其他漏洞可能导致整个 Web 应用被完全控制。',
    '为 /change-password 接口添加 CSRF Token 校验机制。具体修复方案：\n\n'
    '1. 在服务端 /change-password 的 POST 处理逻辑中，在执行密码修改操作之前调用 '
    'validate_csrf_token() 函数进行校验。\n\n'
    '2. 在个人中心页面的密码修改表单中添加隐藏字段 _csrf_token，通过模板引擎的 '
    'csrf_token() 函数生成并填充 Token 值。\n\n'
    '3. CSRF Token 采用每个会话独立生成的强随机字符串（secrets.token_hex(32)），'
    '使用后立即销毁（一次性 Token 机制），防止 Token 被重复利用。',
    '''# 修复前：/change-password 无 CSRF 校验
@app.route("/change-password", methods=["GET", "POST"])
def change_password():
    username = session.get("username")
    if not username:
        return redirect("/login")
    if request.method == "GET":
        current_user_id = USERS.get(username, {}).get("id")
        return redirect(f"/profile?user_id={current_user_id}")

    # 直接从表单接收参数，无 CSRF 校验
    target_username = request.form.get("username", "").strip()
    new_password = request.form.get("new_password", "")
    # ... 直接执行密码修改操作
    USERS[target_username]["password_hash"] = password_hash
    return redirect(f"/profile?user_id={target_user_id}")''',
    '''# 修复后：添加 CSRF Token 校验
@app.route("/change-password", methods=["GET", "POST"])
def change_password():
    username = session.get("username")
    if not username:
        return redirect("/login")
    if request.method == "GET":
        current_user_id = USERS.get(username, {}).get("id")
        return redirect(f"/profile?user_id={current_user_id}")

    # CSRF Token 校验
    if not validate_csrf_token():
        current_user_id = USERS.get(username, {}).get("id")
        return redirect(f"/profile?user_id={current_user_id}")

    target_username = request.form.get("username", "").strip()
    new_password = request.form.get("new_password", "")
    # ... 继续执行密码修改操作
    USERS[target_username]["password_hash"] = password_hash
    return redirect(f"/profile?user_id={target_user_id}")'''
)

doc.add_page_break()

# ── 三、修复措施 ──
add_heading_styled('三、修复措施', level=1)

p = doc.add_paragraph()
run = p.add_run('本次修复涉及以下文件的修改：')
run.font.size = Pt(11)

fixes = [
    ('app.py', '在 /change-password 的 POST 处理逻辑开头添加 validate_csrf_token() 校验调用'),
    ('templates/profile.html', '在密码修改表单中添加 CSRF Token 隐藏字段'),
]
for file, desc in fixes:
    p = doc.add_paragraph(f'  文件：{file}')
    p.runs[0].font.bold = True
    p2 = doc.add_paragraph(f'        修改内容：{desc}')
    p2.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('CSRF 漏洞修复已完成，所有 POST 接口均已启用 CSRF Token 校验机制，'
                '有效防止了跨站请求伪造攻击。')
run.font.size = Pt(10.5)

# ── 保存 ──
output_path = os.path.join(os.path.dirname(__file__), 'day7-CSRF漏洞修复报告.docx')
doc.save(output_path)
print(f'报告已生成：{output_path}')
