#!/usr/bin/env python3
"""生成四川大学本科实习报告"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ===== 辅助函数 =====
def add_centered_text(text, bold=True, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_body_text(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    return p

def add_bold_body(text, indent=True):
    return add_body_text(text, bold=True, indent=indent)

def add_section_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def add_day_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


# ==================== 封面 ====================
for _ in range(4):
    doc.add_paragraph()

add_centered_text("本  科  实  习  报  告", bold=True, size=22)

doc.add_paragraph()
doc.add_paragraph()

# 信息表格
info_table = doc.add_table(rows=7, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ["学    院", "网络空间安全学院", "学生姓名", "郭源"],
    ["专    业", "网络空间安全", "学    号", "202414153"],
    ["年    级", "2024级", "指导教师", "陈腾"],
]

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(3.5)
    row.cells[3].width = Cm(5)

# Merge cells for remaining rows
info_data_full = [
    ["学    院", "", "学生姓名", "郭源"],
    ["专    业", "", "学    号", "202414153"],
    ["年    级", "2024级", "指导教师", "陈腾"],
]

table_data = [
    ("学    院", "网络空间安全学院", "学生姓名", "郭  源"),
    ("专    业", "网络空间安全", "学    号", "202414153"),
    ("年    级", "2024级", "指导教师", "陈  腾"),
]

for i, (c1, c2, c3, c4) in enumerate(table_data):
    for j, val in enumerate([c1, c2, c3, c4]):
        cell = info_table.rows[i].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if j in (0, 2):
            run.bold = True

# 教务处制表行
row4 = info_table.rows[3]
cell40 = row4.cells[0]
cell40.text = ""
cell40_merged = cell40.merge(row4.cells[1])
cell40_merged.text = ""
p = cell40_merged.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("教务处制表")
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

cell42 = row4.cells[2]
cell42_merged = cell42.merge(row4.cells[3])
cell42_merged.text = ""
p = cell42_merged.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("二〇二六年七月二十七日")
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Course info row
row5 = info_table.rows[4]
c = row5.cells[0].merge(row5.cells[1]).merge(row5.cells[2]).merge(row5.cells[3])
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("课程名称：网络安全实践-2        课程号码：314049020        实习周数：2        课程学分：2")
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

row6 = info_table.rows[5]
c = row6.cells[0].merge(row6.cells[1]).merge(row6.cells[2]).merge(row6.cells[3])
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("实习单位：四川讯方信息技术有限公司        实习地点：四川大学空天科学与工程学院四楼401")
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

row7 = info_table.rows[6]
c = row7.cells[0].merge(row7.cells[1]).merge(row7.cells[2]).merge(row7.cells[3])
c.text = ""
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("实习时间：2026年7月18日 - 7月27日")
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 隐藏表格边框
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '  <w:top w:val="{top}" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="{left}" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="{bottom}" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="{right}" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(tcBorders)

for row in info_table.rows:
    for cell in row.cells:
        set_cell_border(cell, top="single", bottom="single", left="single", right="single")

doc.add_page_break()

# ==================== 一、实习目的、要求 ====================
add_section_title("一、实习目的、要求")
doc.add_paragraph()

add_bold_body("1、实习目的")
add_body_text(
    "本次校内综合实训为网络空间安全专业必修实践环节，依托 AI+Web 安全实战模式，旨在让学生掌握 "
    "Python Flask 开发与全套安全工具运维能力，吃透十大 Web 漏洞成因、利用与加固方案。"
    "通过小组完成用户管理平台项目，熟悉 AI 代码审计、自动化攻防脚本开发，建立标准化工程开发流程。"
    "同步锻炼团队协作、项目交付与技术答辩能力，兼顾基础开发与安全攻防双向技能提升，"
    "树立 AI 开发安全风险防范意识，贴合企业网络安全岗位实战需求。"
)

doc.add_paragraph()
add_bold_body("2、实习基本要求")
add_body_text(
    "学生需具备基础编程与网络知识，全程以 AI 辅助开发，重点完成漏洞挖掘与加固，"
    "分层完成基础任务与高阶拓展内容。严格执行每日日清交付、双重考勤制度，"
    "规范 Git 版本管理，完整留存攻防日志与双版本项目代码。"
    "独立完成漏洞实操，杜绝抄袭，缺勤超 2 天将取消考核答辩资格。"
)

doc.add_paragraph()

# ==================== 二、实习主要内容 ====================
add_section_title("二、实习主要内容")
doc.add_paragraph()

add_body_text(
    "本次实习以 Flask 框架为基础，从零开始构建一个完整的用户管理 Web 系统，"
    "并在开发过程中逐步引入十大 Web 安全漏洞，通过「先漏洞后修复」的方式深入理解"
    "每种漏洞的原理、利用方法和修复方案。项目采用 Git 进行版本控制，"
    "代码托管于 GitHub 远程仓库，每日进行任务交付与代码审计。"
    "以下为每日具体工作内容："
)

doc.add_paragraph()

# Day 1
add_day_title("【Day 1】项目初始化与 Flask 基础框架搭建（7月18日）")
add_body_text(
    "第一天主要完成项目的基础环境搭建与 Flask Web 框架学习。使用 Python 3 + Flask 搭建 "
    "Web 应用骨架，创建了基础的目录结构，包括 templates/ 模板目录、static/ 静态资源目录等。"
    "实现了首页路由、登录页面和注册页面的基本框架，学习 Jinja2 模板引擎的用法。"
    "建立了以字典模拟数据库的用户存储结构，包含用户基本信息字段。"
    "配置了 Git 版本管理并完成首次代码提交，建立与 GitHub 远程仓库的连接。"
    "初步了解 MVC 架构模式在 Flask 中的实践方式。"
)

# Day 2
add_day_title("【Day 2】用户认证功能开发与 SQL 注入漏洞引入（7月19日）")
add_body_text(
    "第二天完成用户登录、注册、登出功能的核心逻辑开发。登录模块实现了 Session 会话管理，"
    "用户输入凭据后与存储数据进行比对。注册模块实现了新用户信息录入功能。"
    "在开发过程中，刻意引入了 SQL 注入漏洞——将用户输入直接拼接进 SQL 查询语句，"
    "未使用参数化查询。同时密码采用明文存储，未做任何加密处理。"
    "通过 sqlmap 等工具对 SQL 注入漏洞进行了自动化利用测试，"
    "直观感受 SQL 注入对数据库安全的严重危害。"
    "学习了 Web 安全的基本分类标准和 OWASP Top 10 中常见漏洞类型。"
)

# Day 3
add_day_title("【Day 3】SQL 注入漏洞修复与安全加固（7月20日）")
add_body_text(
    "第三天重点修复了 SQL 注入漏洞。将所有 SQL 查询改为参数化查询（Prepared Statement），"
    "从根本上杜绝了 SQL 注入的可能性。同时完成了以下安全加固工作："
    "引入 Werkzeug 的 generate_password_hash 和 check_password_hash 实现密码 bcrypt 哈希存储；"
    "添加了基于 IP + 用户名的双重登录频率限制（5 次失败锁定 5 分钟）；"
    "配置会话 Cookie 的 HttpOnly 和 SameSite 属性，防止 XSS 劫持 Session；"
    "使用 secrets.token_hex(32) 生成强随机 Secret Key，替代硬编码密钥；"
    "关闭 Flask Debug 模式防止远程代码执行。"
    "产出了 Day3 安全漏洞修复报告（day3漏洞报告.md），记录了漏洞修复前后的代码对比。"
)

# Day 4
add_day_title("【Day 4】文件上传功能开发与上传漏洞修复（7月21日）")
add_body_text(
    "第四天实现了用户头像上传功能，并围绕文件上传进行了 7 层安全加固："
    "扩展名白名单校验——仅允许 .png、.jpg、.gif、.bmp、.webp 格式；"
    "MIME 类型白名单校验——验证 Content-Type 头；"
    "文件魔数签名校验——读取文件头字节验证真实格式；"
    "路径遍历防护——sanitize_filename 函数递归移除 ../ 等路径穿越字符；"
    "条件竞争防护——先保存到临时文件再原子重命名；"
    "空字节截断防护——移除 \\x00 及 URL 编码的 %00；"
    "CVE-2017-15715 换行符绕过防护——移除文件名中的 \\n 和 \\r 字符。"
    "同时将上传目录设置在 static 目录之外，防止直接静态访问，"
    "通过专用路由安全提供文件下载服务。产出了 Day4 文件上传漏洞修复报告。"
)

# Day 5
add_day_title("【Day 5】个人中心、充值功能与越权漏洞修复（7月22日）")
add_body_text(
    "第五天新增了个人中心页面（/profile）和账户充值功能（/recharge）。"
    "个人中心展示用户的详细信息（用户名、邮箱、手机、角色、余额等），"
    "支持通过 user_id 参数查询不同用户。充值功能允许用户为账户增加余额。"
    "在开发中发现了并修复了越权业务逻辑漏洞："
    "IDOR 越权漏洞——普通用户无法通过修改 user_id 参数查看他人资料；"
    "充值越权——充值操作绑定当前会话用户，不从表单接收目标用户名；"
    "负数充值——校验充值金额必须为正数；"
    "CSRF 防护——所有资金操作添加 CSRF Token 校验。"
    "实现了基于角色的访问控制（RBAC），管理员可访问用户管理页面。"
    "产出了 Day5 越权业务逻辑漏洞修复报告。"
)

# Day 6
add_day_title("【Day 6】文件包含漏洞修复与管理员功能开发（7月23日）")
add_body_text(
    "第六天实现了动态页面加载功能（/page 路由），从 pages/ 目录读取 HTML 文件"
    "并渲染到首页。同时完成了管理员用户管理页面（/admin/users），"
    "支持按用户名模糊搜索、列表展示所有用户信息。"
    "在处理文件包含功能时发现了 LFI 漏洞并进行了修复："
    "递归移除路径遍历序列（..、./、\\\\ 等），防止双写绕过；"
    "使用 os.path.realpath 规范化路径并验证文件是否在 pages/ 目录内；"
    "限制了只能读取 pages/ 目录下的文件。"
    "使用 bleach HTML 净化库替换了 {{ content | safe }} 渲染方式，"
    "阻止 XSS 攻击通过恶意 HTML 内容注入。"
    "产出了 Day6 文件包含漏洞修复报告。"
)

# Day 7
add_day_title("【Day 7】CSRF 漏洞修复与密码功能全面加固（7月24日）")
add_body_text(
    "第七天对修改密码功能进行了全面安全审计和修复，共修复 8 个安全漏洞："
    "CSRF-001 无 CSRF 防护——添加 validate_csrf_token() 校验及表单 hidden 字段；"
    "IDOR-001 越权修改他人密码——从 session 获取当前用户名，不从表单接收；"
    "AUTH-001 无需原密码即可修改——增加原密码 check_password_hash 校验；"
    "SESSION-001 修改密码后会话未失效——成功后 session.clear() 强制重新登录；"
    "RATE-001 密码修改无频率限制——复用 LOGIN_LIMIT 机制限制尝试次数；"
    "PWEAK-001 弱密码策略——长度 8 位 + 大小写字母/数字/特殊字符至少三类；"
    "修复了新密码与旧密码相同的问题；"
    "修复了确认密码服务端未验证的问题。"
    "产出了 Day7 CSRF 漏洞修复报告。"
)

# Day 8
add_day_title("【Day 8】SSTI 漏洞修复与新功能开发（7月25日）")
add_body_text(
    "第八天新增了欢迎页（/welcome）和用户反馈（/feedback）两个功能页面。"
    "欢迎页支持通过 URL 参数 name 定制个性化欢迎语；反馈页支持 GET 表单提交 "
    "和 POST 结果展示。在开发中发现并修复了服务器端模板注入（SSTI）漏洞："
    "SSTI-001 /welcome 路由注入——将 f-string 拼接的 render_template_string 改为"
    "模板变量传递方式，阻止通过 URL 参数注入 Jinja2 模板代码；"
    "SSTI-002 /feedback 路由注入——同时修复了 name 和 message 两个注入点。"
    "SSTI 漏洞是 Flask 应用中最危险的高危漏洞之一，攻击者可利用 {{ }} 语法"
    "执行任意 Python 代码，获取服务器控制权。"
    "产出了 Day8 SSTI 漏洞修复报告。"
)

# Day 9
add_day_title("【Day 9】命令执行漏洞修复与全面安全加固（7月26日-7月27日）")
add_body_text(
    "第九天进行了最终的安全加固与工程化改进。新增了 Ping 网络诊断功能（/ping），"
    "允许已登录用户对指定 IP 地址或域名执行网络连通性测试。"
    "在开发该功能时发现了命令注入漏洞（CMD-001）：原实现使用 f-string 拼接用户输入"
    "到系统命令并通过 shell=True 执行，攻击者可提交 8.8.8.8;id 等 payload 执行任意系统命令。"
    "修复措施包括：移除 shell=True 改用参数列表方式调用 subprocess；"
    "增加 IP 地址正则白名单校验和域名格式校验；IP 每段范围 0-255 检查。"
)

add_body_text(
    "随后对项目进行了全面的安全加固增强和工程化改进："
    "添加全局 Content-Security-Policy（CSP）头，形成 XSS 纵深防御；"
    "配置 Session 2 小时自动过期机制，防止会话长期有效；"
    "建立操作审计日志系统，记录登录、密码修改、Ping 执行等全部敏感操作；"
    "实现 Ping 频率限制（每个 IP 每分钟最多 10 次），防止滥用和内网探测；"
    "修复余额浮点数精度问题。"
)

add_body_text(
    "工程化方面完成了五项改进：创建 config.py 集中配置管理模块，"
    "支持环境变量覆盖所有关键配置；编写 requirements.txt 锁定项目依赖版本；"
    "制作 Dockerfile 实现容器化一键部署；"
    "编写 25 个单元测试用例覆盖认证、Ping 注入防护、安全头验证等功能模块；"
    "支持 gunicorn 多 worker 生产级运行模式。"
    "产出了 Day9 命令执行漏洞修复报告。"
)

doc.add_paragraph()

# ==================== 技术栈汇总 ====================
add_section_title("项目技术栈汇总")
doc.add_paragraph()

tech_table = doc.add_table(rows=6, cols=2)
tech_table.style = 'Light Grid Accent 1'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tech_data = [
    ["后端框架", "Flask 3.x + Werkzeug 3.x"],
    ["数据库", "SQLite 3（参数化查询防 SQL 注入）"],
    ["前端模板", "Jinja2 + CSS 渐变风格"],
    ["密码加密", "bcrypt（werkzeug.security）"],
    ["HTML 净化", "bleach 6.x（XSS 防护）"],
    ["版本控制", "Git + GitHub 远程仓库"],
]

for i, (col1, col2) in enumerate(tech_data):
    tech_table.rows[i].cells[0].text = col1
    tech_table.rows[i].cells[1].text = col2
    for cell in [tech_table.rows[i].cells[0], tech_table.rows[i].cells[1]]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ==================== 安全漏洞清单 ====================
add_section_title("已修复安全漏洞清单")
doc.add_paragraph()

vuln_table = doc.add_table(rows=10, cols=4)
vuln_table.style = 'Light Grid Accent 1'
vuln_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["漏洞编号", "漏洞类型", "严重程度", "所属模块"]
vuln_data = [
    ["SQL-001", "SQL 注入", "高危", "登录/注册"],
    ["UPL-001~007", "文件上传绕过", "高危", "头像上传"],
    ["IDOR-001", "越权访问", "高危", "个人中心"],
    ["VULN-002~003", "充值越权/负数充值", "高危", "充值功能"],
    ["LFI-001", "文件包含", "高危", "动态页面"],
    ["XSS-001", "反射型 XSS", "中危", "动态页面"],
    ["CSRF-001", "CSRF 跨站请求伪造", "高危", "改密/充值"],
    ["SSTI-001~002", "模板注入", "严重", "欢迎页/反馈"],
    ["CMD-001", "命令注入", "高危", "Ping 功能"],
]

for i, h in enumerate(headers):
    cell = vuln_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

for i, row_data in enumerate(vuln_data):
    for j, val in enumerate(row_data):
        vuln_table.rows[i+1].cells[j].text = val
        for paragraph in vuln_table.rows[i+1].cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ==================== 三、实习总结 ====================
doc.add_page_break()
add_section_title("三、实习总结")
doc.add_paragraph()

add_body_text(
    "本次为期十天的校内综合实训让我对 Web 安全有了系统而深入的认识。"
    "从 Day1 的项目初始化到 Day9 的全面安全加固，整个实训过程遵循"
    "「先开发后加固，先漏洞后修复」的理念，让我在真实项目环境中深刻理解了"
    "十大 Web 漏洞的成因、利用手法和修复方案。"
)

add_body_text(
    "在技术层面，我系统掌握了 Flask Web 框架的全栈开发能力，包括路由设计、模板渲染、"
    "Session 会话管理、文件上传处理、SQLite 数据库操作等核心技能。"
    "安全方面，我深入理解了 SQL 注入、XSS 跨站脚本、CSRF 跨站请求伪造、"
    "SSTI 模板注入、文件上传绕过、命令注入、路径遍历、越权访问等常见 Web 漏洞的"
    "攻击原理和防御措施。通过亲手修复每一个漏洞，我对安全编码规范有了切身体会。"
)

add_body_text(
    "在工具使用方面，我学会了使用 AI 辅助代码审计和漏洞挖掘，熟练运用 "
    "sqlmap 等安全测试工具进行漏洞验证。同时掌握了 Git 版本控制的最佳实践，"
    "包括分支管理、commit 规范、代码评审流程等工程化开发技能。"
)

add_body_text(
    "在工程思维方面，我深刻认识到安全不是功能的对立面，而是贯穿整个开发生命周期的"
    "必要环节。安全左移——在需求分析和设计阶段就考虑安全威胁，比上线后补救"
    "成本低得多。纵深防御——单一安全措施不足以防御所有攻击，"
    "多层防护（输入校验 + 输出编码 + 安全头 + 日志审计）才能构建可靠的安全体系。"
    "最小权限原则——每个用户只应有完成任务所需的最小权限，"
    "所有资金操作、敏感操作都应绑定当前会话用户。"
)

add_body_text(
    "这次实训还锻炼了我的文档撰写能力和项目交付能力。每天完成安全修复后，"
    "都要撰写结构化的安全修复报告，记录漏洞编号、严重程度、修复前后的代码对比，"
    "这让我的技术文档规范化水平得到了明显提升。"
)

add_body_text(
    "最后，我要感谢陈腾老师的悉心指导和四川讯方信息技术有限公司提供的实训平台支持。"
    "这次实训让我从一个只会写简单脚本的新手，成长为具备安全意识的 Web 开发者。"
    "今后我将继续保持对网络安全的学习热情，将本次实训中习得的安全编码习惯"
    "应用到未来的开发工作中，为国家的网络安全事业贡献自己的力量。"
)

doc.add_paragraph()
doc.add_paragraph()

# ==================== 成绩评定页 ====================
add_section_title("实习成绩评定")
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("指导教师签名：                          ")
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("年    月    日")
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 保存 =====
output_path = "/opt/Class01/四川大学本科实习报告-郭源-202414153.docx"
doc.save(output_path)
print(f"实习报告已生成: {output_path}")
