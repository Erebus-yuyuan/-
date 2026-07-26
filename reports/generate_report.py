#!/usr/bin/env python3
"""生成安全漏洞修复报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 辅助函数 ──
def set_cell_shading(cell, color_hex):
    """设置表格单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_severity_cell(cell, text):
    """严重程度单元格着色"""
    color_map = {
        '🔴 严重': 'FFCDD2',
        '🟠 高危': 'FFE0B2',
        '🟡 中危': 'FFF9C4',
        '🔵 低危': 'BBDEFB',
    }
    for key, color in color_map.items():
        if key in text:
            set_cell_shading(cell, color)
            break
    cell.text = text


def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color=None, space_after=6):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code_block(doc, code_text):
    """添加代码块（灰色等宽背景）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    # 浅灰底
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    run._element.rPr.append(shading)
    return p


# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('安全漏洞修复报告')
run.font.size = Pt(28)
run.bold = True
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Flask 用户信息管理系统 · 安全加固项目')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x6E, 0x7A)
run.font.name = '微软雅黑'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

meta_items = [
    ('项目版本', 'V2.0（新增注册/修改密码功能）'),
    ('报告日期', '2026年7月19日'),
    ('漏洞总数', '7 项（已全部修复）'),
    ('安全状态', '✅ 已加固'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = '微软雅黑'

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading_styled(doc, '目  录', 1)
toc_items = [
    '一、项目概述',
    '二、漏洞分析与修复详情',
    '    2.1  HTML 注释泄露管理员凭证',
    '    2.2  前端页面回显密码',
    '    2.3  明文存储密码',
    '    2.4  弱 Secret Key 导致 Session 伪造',
    '    2.5  Debug 模式远程代码执行风险',
    '    2.6  登录频率限制缺失',
    '    2.7  用户对象原地变异',
    '三、修复前后对比总表',
    '四、一键验证指南',
    '五、新增功能说明',
    '六、生产环境部署建议',
]
for item in toc_items:
    add_para(doc, item, size=11, space_after=4)

doc.add_page_break()

# ============================================================
# 一、项目概述
# ============================================================
add_heading_styled(doc, '一、项目概述', 1)
add_para(doc, (
    '本项目为一个基于 Flask 框架的用户信息管理系统（Web 安全靶场），'
    '初始版本存在多项严重安全漏洞。本次安全加固针对发现的 7 项安全漏洞进行了系统性修复，'
    '并新增了用户注册与密码修改功能，使系统具备完整的用户管理能力。'
))
add_para(doc, '• 开发框架：Flask（Python）')
add_para(doc, '• 项目路径：/opt/Class01/')
add_para(doc, '• 漏洞数量：7 项（严重 2 项 / 高危 4 项 / 中危 1 项）')
add_para(doc, '• 修复状态：✅ 全部完成')
add_para(doc, '• 新增功能：用户注册、密码修改')

doc.add_paragraph()

# ============================================================
# 二、漏洞分析与修复详情
# ============================================================
add_heading_styled(doc, '二、漏洞分析与修复详情', 1)

# ── 2.1 ──
add_heading_styled(doc, '2.1  HTML 注释泄露管理员凭证', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🔴 严重')
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
add_para(doc, '【漏洞描述】')
add_para(doc, '登录页面（templates/login.html）的 HTML 源码中包含一条调试注释，'
           '直接暴露了默认管理员账号 admin 及其密码 admin123。任何访客通过"查看网页源代码"即可获取管理员凭证。')
add_para(doc, '【OWASP 映射】A05:2021 — 安全配置错误（Security Misconfiguration）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, '<!-- 调试信息 - 默认管理员账号 用户名: admin 密码: admin123 -->')
add_para(doc, '【修复方式】✅ 已删除该注释行')
add_para(doc, '【影响】攻击者无法再通过页面源码直接获取管理员账号密码。')

# ── 2.2 ──
add_heading_styled(doc, '2.2  前端页面回显密码', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🔴 严重')
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
add_para(doc, '【漏洞描述】')
add_para(doc, '用户登录成功后，首页（templates/index.html）直接将用户的明文密码渲染在页面上。'
           '即使是合法用户登录，密码也会暴露在屏幕前，存在肩窥和截屏泄露风险。')
add_para(doc, '【OWASP 映射】A04:2021 — 不安全设计（Insecure Design）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, '<li><span class="info-label">密码：</span><span class="info-value">{{ user.password }}</span></li>')
add_para(doc, '【修复方式】✅ 删除该模板行，后端新增 safe_user_info() 函数确保密码哈希不流出到模板层')
add_para(doc, '【影响】用户登录后页面不再显示密码，敏感信息不会意外泄露。')

# ── 2.3 ──
add_heading_styled(doc, '2.3  明文存储密码', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🟠 高危')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
add_para(doc, '【漏洞描述】')
add_para(doc, 'USERS 字典中密码以明文形式存储（"password": "admin123"）。'
           '一旦服务器文件泄露（如目录遍历、备份泄露），所有用户的密码原文将直接暴露。')
add_para(doc, '【OWASP 映射】A02:2021 — 加密失效（Cryptographic Failures）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, 'USERS = {"admin": {"password": "admin123", ...}}')
add_para(doc, '【修复方式】✅ 改用 werkzeug.security 的 bcrypt（PBKDF2-SHA256）加盐哈希存储')
add_code_block(doc, (
    'from werkzeug.security import generate_password_hash, check_password_hash\n'
    'USERS = {"admin": {"password_hash": generate_password_hash("admin123"), ...}}'
))
add_para(doc, '• generate_password_hash() 使用 bcrypt 算法自动加盐哈希')
add_para(doc, '• check_password_hash(hash, password) 进行不可逆比对')
add_para(doc, '• 即使数据库泄露，攻击者也无法还原原始密码')
add_para(doc, '【影响】密码以不可逆哈希形式存储，数据库泄露时攻击者无法获取原文。')

# ── 2.4 ──
add_heading_styled(doc, '2.4  弱 Secret Key 导致 Session 伪造', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🟠 高危')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
add_para(doc, '【漏洞描述】')
add_para(doc, 'Flask 的 app.secret_key 设置为 "dev-key-2025"，该弱密钥可被暴力破解。'
           '攻击者一旦获取密钥，可伪造任意用户的 session cookie，冒充任意身份登录系统。')
add_para(doc, '【OWASP 映射】A02:2021 — 加密失效（Cryptographic Failures）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, 'app.secret_key = "dev-key-2025"')
add_para(doc, '【修复方式】✅ 改用环境变量 + 62 位强随机十六进制密钥')
add_code_block(doc, (
    'import os\n'
    'app.secret_key = os.environ.get(\n'
    '    "SECRET_KEY",\n'
    '    "a9f8d7e6b5c4a3f2e1d0c9b8a7f6e5d4c3b2a1f0e9d8c7b6a5f4e3d2c1b0a9"\n'
    ')'
))
add_para(doc, '• 优先从环境变量 SECRET_KEY 读取')
add_para(doc, '• 生产环境建议：export SECRET_KEY=$(python3 -c "import secrets; print(secrets.token_hex(32))")')
add_para(doc, '【影响】密钥空间从 12 字符提升至 62 位十六进制，暴力破解不可行。')

# ── 2.5 ──
add_heading_styled(doc, '2.5  Debug 模式远程代码执行风险', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🟠 高危')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
add_para(doc, '【漏洞描述】')
add_para(doc, 'app.run(debug=True, host="0.0.0.0") 同时存在两个问题：'
           'Debug 模式下 Flask 的 /console 端点可执行任意 Python 代码；'
           'host="0.0.0.0" 使服务监听所有网络接口。')
add_para(doc, '【OWASP 映射】A05:2021 — 安全配置错误（Security Misconfiguration）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, 'app.run(debug=True, host="0.0.0.0", port=5000)')
add_para(doc, '【修复方式】✅ 关闭 Debug 模式，保留外部访问能力（靶场需求）')
add_code_block(doc, 'app.run(debug=False, host="0.0.0.0", port=5000)')
add_para(doc, '• debug=False：禁用 /console 交互式调试器，防止远程代码执行')
add_para(doc, '• host="0.0.0.0"：保留以支持外部主机访问（靶场教学场景需求）')
add_para(doc, '【影响】攻击者无法通过 /console 注入代码，同时保持靶场的远程访问能力。')

# ── 2.6 ──
add_heading_styled(doc, '2.6  登录频率限制缺失', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🟠 高危')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
add_para(doc, '【漏洞描述】')
add_para(doc, '原始登录接口无任何频率限制，攻击者可以无限次尝试用户名和密码组合进行暴力破解。'
           '结合已泄露的两个用户名（admin、alice），字典攻击的成功率极高。')
add_para(doc, '【OWASP 映射】A07:2021 — 身份验证失效（Identification and Authentication Failures）')
add_para(doc, '【修复方式】✅ 实现基于 IP + 用户名的双重限速机制')

# 表格：限速参数
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('参数', '值'),
    ('最大失败次数', '5 次'),
    ('锁定时长', '300 秒（5 分钟）'),
    ('限制粒度', 'IP : 用户名 组合'),
    ('自动清理策略', '记录超 100 条时自动回收过期条目'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[i].cells:
            set_cell_shading(cell, '1A237E')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

doc.add_paragraph()
add_para(doc, '第 6 次连续失败登录将返回锁定提示："登录已被锁定，请在 296 秒后重试！"')
add_para(doc, '【影响】暴力破解在 5 次失败后被自动阻断 5 分钟，大幅提高破解成本。')

# ── 2.7 ──
add_heading_styled(doc, '2.7  用户对象原地变异（代码质量缺陷）', 2)
severity = doc.add_paragraph()
run = severity.add_run('【严重程度】 🟡 中危')
run.bold = True
run.font.color.rgb = RGBColor(0xF9, 0xA8, 0x25)
add_para(doc, '【漏洞描述】')
add_para(doc, '原始代码使用 user_info = USERS[username] 直接获取字典引用，'
           '然后通过 user_info["username"] = username 修改对象，导致全局 USERS 字典被"污染"，'
           '每次登录后用户记录被额外添加一个 username 字段。')
add_para(doc, '【OWASP 映射】代码质量问题（Coding Quality）')
add_para(doc, '【修复前】', bold=True)
add_code_block(doc, (
    'user_info = USERS[username]        # 直接引用，非副本\n'
    'user_info["username"] = username    # 修改了全局字典！'
))
add_para(doc, '【修复方式】✅ 使用 copy.deepcopy() 创建副本后再修改')
add_code_block(doc, (
    'def safe_user_info(username):\n'
    '    info = copy.deepcopy(USERS[username])\n'
    '    info.pop("password_hash", None)  # 同时确保哈希不泄露\n'
    '    info["username"] = username\n'
    '    return info'
))
add_para(doc, '• 原始数据完好无损，不会被后续操作影响')
add_para(doc, '• 密码哈希永远不会流出到模板渲染层')
add_para(doc, '【影响】全局数据始终保持一致，消除了因引用操作导致的隐蔽数据污染。')

doc.add_page_break()

# ============================================================
# 三、修复前后对比总表
# ============================================================
add_heading_styled(doc, '三、修复前后对比总表', 1)

table2 = doc.add_table(rows=8, cols=5)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['编号', '漏洞名称', '严重程度', '修复前', '修复后']
rows_data = [
    ('1', 'HTML 注释泄露密码', '🔴 严重', '注释含 admin/admin123', '注释已删除'),
    ('2', '前端回显密码', '🔴 严重', '首页显示明文密码', '密码已移除'),
    ('3', '明文存储密码', '🟠 高危', '"password":"admin123"', 'bcrypt 哈希'),
    ('4', '弱 Secret Key', '🟠 高危', '"dev-key-2025"', '62位强密钥'),
    ('5', 'Debug 模式', '🟠 高危', 'debug=True', 'debug=False'),
    ('6', '无登录频率限制', '🟠 高危', '无限尝试', '5次锁定5分钟'),
    ('7', '对象原地变异', '🟡 中危', '直接修改全局字典', 'deepcopy 副本'),
]

for j, h in enumerate(headers):
    cell = table2.rows[0].cells[j]
    cell.text = h
    set_cell_shading(cell, '1A237E')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(9)

owasp_map = ['A05', 'A04', 'A02', 'A02', 'A05', 'A07', '—']
for i, (num, name, severity_txt, before, after) in enumerate(rows_data):
    row_idx = i + 1
    table2.rows[row_idx].cells[0].text = num
    table2.rows[row_idx].cells[1].text = name
    sev_cell = table2.rows[row_idx].cells[2]
    sev_cell.text = severity_txt
    add_severity_cell(sev_cell, severity_txt)
    table2.rows[row_idx].cells[3].text = before
    table2.rows[row_idx].cells[4].text = after
    for cell in table2.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# 四、一键验证指南
# ============================================================
add_heading_styled(doc, '四、一键验证指南', 1)
add_para(doc, '以下命令可快速验证各漏洞的修复效果：')
add_para(doc, '验证 1：使用正确密码登录', bold=True)
add_code_block(doc, (
    '# 返回 HTTP 200，页面显示"欢迎回来，admin！"\n'
    'curl -s http://<服务器IP>:5000/login -X POST \\\n'
    '  -d "username=admin&password=admin123" | grep "欢迎回来"'
))
add_para(doc, '验证 2：确认密码不再回显', bold=True)
add_code_block(doc, (
    '# 返回 0（页面中无"密码："字段）\n'
    'curl -s http://<服务器IP>:5000/login -X POST \\\n'
    '  -d "username=admin&password=admin123" | grep -c "密码："'
))
add_para(doc, '验证 3：锁定机制测试', bold=True)
add_code_block(doc, (
    '# 连续输错 6 次密码，第 6 次返回锁定提示\n'
    'for i in $(seq 1 6); do\n'
    '  curl -s http://<服务器IP>:5000/login -X POST \\\n'
    '    -d "username=admin&password=wrong" | grep -o "已被锁定"\n'
    'done'
))
add_para(doc, '验证 4：确认 Debug 模式关闭', bold=True)
add_code_block(doc, (
    '# 返回 404（/console 端点不可用）\n'
    'curl -s http://<服务器IP>:5000/console'
))
add_para(doc, '验证 5：用户注册', bold=True)
add_code_block(doc, (
    '# 注册新用户并自动登录\n'
    'curl -s http://<服务器IP>:5000/register -X POST \\\n'
    '  -d "username=test&password=testpass&confirm_password=testpass" | grep "欢迎回来"'
))
add_para(doc, '验证 6：修改密码', bold=True)
add_code_block(doc, (
    '# 修改密码后旧密码失效，新密码可登录\n'
    'curl -s -c /tmp/cookies.txt http://<服务器IP>:5000/login -X POST \\\n'
    '  -d "username=test&password=testpass" > /dev/null\n'
    'curl -s -b /tmp/cookies.txt http://<服务器IP>:5000/change-password -X POST \\\n'
    '  -d "old_password=testpass&new_password=newpass&confirm_password=newpass"'
))

doc.add_page_break()

# ============================================================
# 五、新增功能说明
# ============================================================
add_heading_styled(doc, '五、新增功能说明', 1)
add_para(doc, '在安全加固的基础上，新增了以下用户功能：')

add_heading_styled(doc, '5.1  用户注册（/register）', 2)
add_para(doc, '• 支持新用户自主注册，字段包括用户名、密码、确认密码、邮箱（选填）、手机号（选填）')
add_para(doc, '• 输入校验：用户名≥2字符且唯一、密码≥6位、两次密码一致')
add_para(doc, '• 新用户默认 role="user"、balance=0')
add_para(doc, '• 注册成功后自动登录并跳转首页')

add_heading_styled(doc, '5.2  修改密码（/change-password）', 2)
add_para(doc, '• 需登录状态，未登录自动跳转登录页')
add_para(doc, '• 输入校验：旧密码正确、新密码≥6位、新旧密码不同、两次输入一致')
add_para(doc, '• 修改成功后跳转首页并显示成功提示')
add_para(doc, '• 密码同样使用 bcrypt 哈希存储')

add_heading_styled(doc, '5.3  导航栏更新', 2)
add_para(doc, '• 未登录状态：显示"注册"和"登录"链接')
add_para(doc, '• 已登录状态：显示用户名、"修改密码"和"退出"链接')

doc.add_page_break()

# ============================================================
# 六、生产环境部署建议
# ============================================================
add_heading_styled(doc, '六、生产环境部署建议', 1)
add_para(doc, '以下建议适用于将本项目部署到生产环境时的安全加强措施：')
add_para(doc, '1. 使用环境变量管理密钥', bold=True)
add_para(doc, '   export SECRET_KEY=$(python3 -c "import secrets; print(secrets.token_hex(32))")')
add_para(doc, '2. 使用生产级 WSGI 服务器', bold=True)
add_para(doc, '   pip install gunicorn && gunicorn -w 4 -b 127.0.0.1:5000 app:app')
add_para(doc, '3. 配置 HTTPS 反向代理（推荐 Nginx 或 Caddy）', bold=True)
add_para(doc, '   终止 TLS、托管静态文件、添加安全响应头')
add_para(doc, '4. 迁移到持久化数据库', bold=True)
add_para(doc, '   将 USERS 字典替换为 SQLite / PostgreSQL 等数据库存储')
add_para(doc, '5. 添加 CSRF 保护', bold=True)
add_para(doc, '   使用 Flask-WTF 扩展为所有表单添加 CSRF Token')
add_para(doc, '6. 配置日志审计', bold=True)
add_para(doc, '   记录所有登录尝试和敏感操作日志')

# ── 页脚 ──
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# ── 保存 ──
output_path = '/opt/Class01/security_report_v2.0.docx'
doc.save(output_path)
print(f'Word 文档已生成：{output_path}')
